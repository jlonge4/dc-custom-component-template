from typing import Any, Dict, List, Optional, Union
from pathlib import Path

from haystack import Document, component
import io
import logging
from haystack.dataclasses import ByteStream
from haystack.lazy_imports import LazyImport
from haystack.components.converters.utils import (
    get_bytestream_from_source,
    normalize_metadata,
)

logger = logging.getLogger(__name__)

with LazyImport("Run 'pip install python-docx'") as docx_import:
    import docx


@component
class DOCXTablesToDocument:
    """
    A component for extracting tables from DOCX files and converting them to Markdown-formatted Document objects.

    This component processes DOCX files, extracts all tables from each file, and converts each table
    into a Markdown-formatted string. Each table is then created as a separate Document object with
    appropriate metadata, including the page number where the table was found.

    ### Usage example

    ```python
    from haystack import Pipeline
    from haystack.dataclasses import ByteStream
    from custom_components import DOCXTablesToDocument

    # Create an instance of the component
    docx_tables_converter = DOCXTablesToDocument()

    # Create a pipeline with the component
    pipeline = Pipeline()
    pipeline.add_component("table_extractor", docx_tables_converter)

    # Prepare your DOCX file as a ByteStream
    with open("path/to/your/document.docx", "rb") as file:
        docx_bytestream = ByteStream(file.read())

    # Run the pipeline
    result = pipeline.run(
        {
            "table_extractor": {
                "sources": [docx_bytestream],
                "meta": {"filename": "document.docx"}
            }
        }
    )

    # Access the extracted tables as Document objects
    documents = result["table_extractor"]["documents"]
    for doc in documents:
        print(f"Table from page {doc.meta['page_num']}:")
        print(doc.content)
        print("---")
    """

    def __init__(self):
        """
        Create a DOCXTablesToDocument component.
        """
        docx_import.check()

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path, ByteStream]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        """
        Split documents into smaller parts using character lengths.

        :param documents: The documents to split.

        :returns: A dictionary with the following key:
            - `documents`: List of documents with the split texts. Each document includes:
                - A metadata field `source_id` to track the original document.
                - A metadata field `source_offsets` to track the original character offsets.
                - A metadata field `source_part` to track the part number of the split text.
                - All other metadata copied from the original document.
        """
        documents = []
        meta_list = normalize_metadata(meta=meta, sources_count=len(sources))

        for source, metadata in zip(sources, meta_list):
            try:
                bytestream = get_bytestream_from_source(source)
            except Exception as e:
                logger.warning(
                    "Could not read {source}. Skipping it. Error: {error}",
                    source=source,
                    error=e,
                )
                continue
            try:
                file = docx.Document(io.BytesIO(bytestream.data))
            except Exception as e:
                logger.warning(
                    "Could not read {source} and convert it to a DOCX Document, skipping. Error: {error}",
                    source=source,
                    error=e,
                )
                continue
            for page_num, table in enumerate(file.tables, start=1):
                mkdown = []
                try:
                    for i, row in enumerate(table.rows):
                        md_row = [cell.text.strip() for cell in row.cells]
                        mkdown.append("|" + "|".join(md_row) + "|")
                        if i == 0:
                            mkdown.append(
                                "|" + "|".join(["---" for _ in row.cells]) + "|"
                            )
                    mkdown = "\n".join(mkdown)
                    metadata["page_num"] = page_num
                    merged_metadata = {**bytestream.meta, **metadata}
                    documents.append(
                        Document(content=str(mkdown), meta=merged_metadata)
                    )
                except Exception as e:
                    logger.error(f"Error: {e}. Skipping table on page {page_num}")
                    continue
            

        return {"documents": documents}
